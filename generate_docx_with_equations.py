from __future__ import annotations

import argparse
import json
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


DEFAULT_XSL_CANDIDATES = [
    Path(r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
]


def find_mml2omml_xsl(explicit_path: str | None) -> Path:
    if explicit_path:
        path = Path(explicit_path)
        if path.is_file():
            return path
        raise FileNotFoundError(f"找不到指定的 MML2OMML.XSL: {path}")

    for candidate in DEFAULT_XSL_CANDIDATES:
        if candidate.is_file():
            return candidate

    raise FileNotFoundError(
        "未找到 MML2OMML.XSL。请安装 Microsoft Word，或用 --xsl 指定该文件路径。"
    )


def set_run_font(run, font_name: str, font_size: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name: str, font_size: int) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = tbl_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_element)
        edge_element.set(qn("w:val"), "nil")


class EquationDocBuilder:
    def __init__(self, xsl_path: Path) -> None:
        self.transform = etree.XSLT(etree.parse(str(xsl_path)))

    def latex_to_omml(self, latex: str):
        mathml = latex_to_mathml(latex)
        mathml_root = etree.fromstring(mathml.encode("utf-8"))
        omml_tree = self.transform(mathml_root)
        return deepcopy(omml_tree.getroot())

    def append_formula_run(self, paragraph, latex: str) -> None:
        run = paragraph.add_run()
        run._element.append(self.latex_to_omml(latex))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def format_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, font_name="黑体", font_size=16, bold=False)


def add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, font_name="宋体", font_size=12)


def add_mixed_paragraph(document: Document, builder: EquationDocBuilder, parts: list[dict[str, str]]) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)

    for part in parts:
        text = part.get("text")
        latex = part.get("latex")

        if text is not None:
            run = paragraph.add_run(text)
            set_run_font(run, font_name="宋体", font_size=12)
        elif latex is not None:
            builder.append_formula_run(paragraph, latex)
        else:
            raise ValueError(f"不支持的 mixed_paragraph 片段: {part}")


def add_equation_block(document: Document, builder: EquationDocBuilder, latex: str, number: str) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    widths = [Cm(1.5), Cm(11.7), Cm(2.2)]
    row = table.rows[0]
    row.cells[0].width = widths[0]
    row.cells[1].width = widths[1]
    row.cells[2].width = widths[2]

    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    center_paragraph = row.cells[1].paragraphs[0]
    center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.append_formula_run(center_paragraph, latex)

    number_paragraph = row.cells[2].paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_run = number_paragraph.add_run(number)
    set_run_font(number_run, font_name="Times New Roman", font_size=12)

    document.add_paragraph()


def build_document(config: dict[str, Any], xsl_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    builder = EquationDocBuilder(xsl_path)

    for block in config["blocks"]:
        block_type = block["type"]
        if block_type == "heading":
            add_heading(document, block["text"])
        elif block_type == "paragraph":
            add_text_paragraph(document, block["text"])
        elif block_type == "mixed_paragraph":
            add_mixed_paragraph(document, builder, block["parts"])
        elif block_type == "equation":
            add_equation_block(document, builder, block["latex"], block["number"])
        else:
            raise ValueError(f"不支持的 block 类型: {block_type}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="将结构化文字 + LaTeX 公式排版成带可编辑公式的 Word 文档。"
    )
    parser.add_argument("input", help="输入 JSON 文件路径")
    parser.add_argument("output", help="输出 DOCX 文件路径")
    parser.add_argument(
        "--xsl",
        help="MML2OMML.XSL 路径；不传时会在常见的 Microsoft Office 安装目录中自动查找。",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)
    xsl_path = find_mml2omml_xsl(args.xsl)

    config = json.loads(input_path.read_text(encoding="utf-8"))
    build_document(config, xsl_path, output_path)
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    main()
