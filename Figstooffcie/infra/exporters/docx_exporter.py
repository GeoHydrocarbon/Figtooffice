from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from infra.equation.omml_converter import OmmlConverter


INLINE_MATH = re.compile(r"\$(?!\$)(.+?)\$(?!\$)", re.DOTALL)
HEADING_LINE = re.compile(r"^(#{1,6})\s+(.+)$")


def export_markdown_to_docx(markdown: str, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    converter = OmmlConverter()

    for kind, payload in iter_markdown_events(markdown):
        if kind == "heading":
            level, text = payload
            if level == 1:
                add_heading(document, text)
            else:
                document.add_heading(text, level=min(max(level - 1, 0), 9))
        elif kind == "paragraph":
            add_paragraph_markdown(document, converter, payload)
        elif kind == "display":
            add_centered_display_equation(document, converter, payload)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_run_font(run, font_name: str, font_size: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def format_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, font_name="黑体", font_size=16)


def add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, font_name="宋体", font_size=12)


def add_mixed_paragraph(document: Document, converter: OmmlConverter, parts: list[dict[str, str]]) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    for part in parts:
        if "text" in part:
            run = paragraph.add_run(part["text"])
            set_run_font(run, font_name="宋体", font_size=12)
        elif "latex" in part:
            run = paragraph.add_run()
            run._element.append(converter.to_omml(part["latex"]))


def add_centered_display_equation(document: Document, converter: OmmlConverter, latex: str) -> None:
    if not latex.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run._element.append(converter.to_omml(latex))


def add_paragraph_markdown(document: Document, converter: OmmlConverter, text: str) -> None:
    parts = text_to_mixed_parts(text)
    if len(parts) == 1 and "text" in parts[0]:
        add_text_paragraph(document, parts[0]["text"])
        return
    add_mixed_paragraph(document, converter, parts)


def text_to_mixed_parts(text: str) -> list[dict[str, str]]:
    parts: list[dict[str, str]] = []
    pos = 0
    for match in INLINE_MATH.finditer(text):
        if match.start() > pos:
            chunk = text[pos : match.start()]
            if chunk:
                parts.append({"text": chunk})
        latex = match.group(1).strip()
        if latex:
            parts.append({"latex": latex})
        pos = match.end()
    if pos < len(text):
        tail = text[pos:]
        if tail:
            parts.append({"text": tail})
    if not parts:
        parts.append({"text": text})
    return parts


def iter_markdown_events(markdown: str):
    lines = markdown.splitlines()
    index = 0
    total = len(lines)
    while index < total:
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        heading_match = HEADING_LINE.match(stripped)
        if heading_match:
            yield "heading", (len(heading_match.group(1)), heading_match.group(2).strip())
            index += 1
            continue

        if stripped.startswith("$$"):
            if stripped.endswith("$$") and stripped.count("$$") == 2 and len(stripped) >= 6:
                yield "display", stripped[2:-2].strip()
                index += 1
                continue

            inner: list[str] = []
            if stripped != "$$":
                remainder = stripped[2:].strip()
                if remainder.endswith("$$"):
                    yield "display", remainder[:-2].strip()
                    index += 1
                    continue
                inner.append(remainder)
            index += 1
            while index < total:
                line = lines[index]
                candidate = line.strip()
                if candidate.endswith("$$"):
                    prefix = line.rsplit("$$", 1)[0]
                    if prefix.strip():
                        inner.append(prefix.rstrip())
                    index += 1
                    break
                inner.append(line)
                index += 1
            else:
                raise ValueError("存在未闭合的块级公式。")
            yield "display", "\n".join(inner).strip()
            continue

        buffer = [lines[index]]
        index += 1
        while index < total:
            line = lines[index]
            stripped_line = line.strip()
            if not stripped_line or stripped_line.startswith("#") or stripped_line.startswith("$$"):
                break
            buffer.append(line)
            index += 1
        paragraph = "\n".join(buffer).strip()
        if paragraph:
            yield "paragraph", paragraph
