from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from generate_docx_with_equations import (
    EquationDocBuilder,
    add_heading,
    add_mixed_paragraph,
    add_text_paragraph,
    configure_document,
)


INLINE_MATH = re.compile(r"\$(?!\$)(.+?)\$(?!\$)", re.DOTALL)
HEADING_LINE = re.compile(r"^(#{1,6})\s+(.+)$")


def text_to_mixed_parts(text: str) -> list[dict[str, str]]:
    """将含 $...$ 行内公式的文本拆成 generate_docx_with_equations 所需的 parts 列表。"""
    parts: list[dict[str, str]] = []
    pos = 0
    for m in INLINE_MATH.finditer(text):
        if m.start() > pos:
            chunk = text[pos : m.start()]
            if chunk:
                parts.append({"text": chunk})
        latex = m.group(1).strip()
        if latex:
            parts.append({"latex": latex})
        pos = m.end()
    if pos < len(text):
        tail = text[pos:]
        if tail:
            parts.append({"text": tail})
    if not parts:
        parts.append({"text": text})
    return parts


def add_centered_display_equation(document: Document, builder: EquationDocBuilder, latex: str) -> None:
    """独立公式：居中一段，无编号、无三列表格。"""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.append_formula_run(paragraph, latex)


def add_paragraph_markdown(document: Document, builder: EquationDocBuilder, text: str) -> None:
    parts = text_to_mixed_parts(text)
    if len(parts) == 1 and "text" in parts[0]:
        add_text_paragraph(document, parts[0]["text"])
    else:
        add_mixed_paragraph(document, builder, parts)


def iter_markdown_events(md: str):
    """将 Markdown 拆成事件流：heading / display / paragraph。"""
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        s = lines[i].strip()
        if s == "":
            i += 1
            continue

        hm = HEADING_LINE.match(s)
        if hm:
            yield "heading", (len(hm.group(1)), hm.group(2).strip())
            i += 1
            continue

        if s.startswith("$$"):
            if s.endswith("$$") and len(s) >= 6 and s.count("$$") == 2:
                yield "display", s[2:-2].strip()
                i += 1
                continue

            inner: list[str] = []
            if s != "$$":
                rest = s[2:].strip()
                if rest.endswith("$$"):
                    yield "display", rest[:-2].strip()
                    i += 1
                    continue
                inner.append(rest)
            i += 1
            while i < n:
                ln = lines[i]
                st = ln.strip()
                if st.endswith("$$"):
                    pre = ln.rsplit("$$", 1)[0]
                    if pre.strip():
                        inner.append(pre.rstrip())
                    i += 1
                    break
                inner.append(ln)
                i += 1
            else:
                raise ValueError("未闭合的 $$ 块（缺少结尾 $$）")
            yield "display", "\n".join(inner).strip()
            continue

        buf = [lines[i]]
        i += 1
        while i < n:
            t = lines[i]
            ts = t.strip()
            if ts == "":
                break
            if ts.startswith("#") or ts.startswith("$$"):
                break
            buf.append(lines[i])
            i += 1
        para = "\n".join(buf).strip()
        if para:
            yield "paragraph", para


def build_docx_from_markdown(markdown: str, xsl_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    builder = EquationDocBuilder(xsl_path)

    for kind, payload in iter_markdown_events(markdown):
        if kind == "heading":
            level, text = payload
            if level == 1:
                add_heading(document, text)
            else:
                document.add_heading(text, level=min(max(level - 1, 0), 9))
        elif kind == "display":
            latex: str = payload
            if latex.strip():
                add_centered_display_equation(document, builder, latex)
        elif kind == "paragraph":
            add_paragraph_markdown(document, builder, payload)
        else:
            raise ValueError(f"未知事件: {kind}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
